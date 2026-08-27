import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_short_document():
    doc = docx.Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    NAVY = RGBColor(15, 44, 89)
    SLATE = RGBColor(71, 85, 105)
    DARK_TEXT = RGBColor(30, 41, 59)

    # Header
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_header.add_run("COER UNIVERSITY\n")
    run_inst.font.name = "Arial"
    run_inst.font.size = Pt(15)
    run_inst.font.bold = True
    run_inst.font.color.rgb = NAVY

    run_sub = p_header.add_run("AI TASK FORCE — FACE RECOGNITION ATTENDANCE SYSTEM\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(10.5)
    run_sub.font.bold = True
    run_sub.font.color.rgb = SLATE

    run_divider = p_header.add_run("―" * 45 + "\n")
    run_divider.font.color.rgb = SLATE

    # Memo Table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    meta_data = [
        ("TO:", "Higher Authority, COER University"),
        ("FROM:", "AI Task Force"),
        ("DATE:", "August 25, 2026"),
        ("SUBJECT:", "Urgent: Project Delay Due to Photo Submissions Non-Compliance by Colleges")
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row = table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = Inches(1.1)
        cell_val.width = Inches(5.4)
        
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.name = "Arial"
        r_lbl.font.size = Pt(10)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = NAVY
        
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.name = "Arial"
        r_val.font.size = Pt(10)
        if label == "SUBJECT:":
            r_val.font.bold = True
            r_val.font.color.rgb = NAVY
        else:
            r_val.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    # Salutation
    p_sal = doc.add_paragraph()
    r_sal = p_sal.add_run("Respected Sir/Madam,")
    r_sal.font.name = "Arial"
    r_sal.font.size = Pt(10.5)
    r_sal.font.bold = True
    r_sal.font.color.rgb = DARK_TEXT

    # Brief Body
    p_body = doc.add_paragraph()
    r_body = p_body.add_run(
        "The completion and university-wide deployment of the Face Recognition Attendance System is getting "
        "significantly delayed because colleges are failing to provide proper photo datasets despite repeated notices."
    )
    r_body.font.name = "Arial"
    r_body.font.size = Pt(10.5)
    r_body.font.color.rgb = DARK_TEXT
    p_body.paragraph_format.line_spacing = 1.15
    p_body.paragraph_format.space_after = Pt(8)

    # Issues Section Header
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("Key Issues:")
    r_h1.font.name = "Arial"
    r_h1.font.size = Pt(11)
    r_h1.font.bold = True
    r_h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(4)

    # Single-line issues
    single_line_issues = [
        "Colleges are providing improper and low-quality photos unsuitable for face recognition.",
        "Colleges are still not sharing required photos despite multiple iterative notices and communications."
    ]

    for issue in single_line_issues:
        p_iss = doc.add_paragraph(style='List Bullet')
        p_iss.paragraph_format.line_spacing = 1.15
        p_iss.paragraph_format.space_after = Pt(4)
        r_iss = p_iss.add_run(issue)
        r_iss.font.name = "Arial"
        r_iss.font.size = Pt(10.5)
        r_iss.font.color.rgb = DARK_TEXT

    # Request Action
    p_act = doc.add_paragraph()
    r_act = p_act.add_run(
        "\nDue to these issues, the project is taking unnecessary extra time. "
        "We kindly request your urgent directive to all College Heads to submit compliant photos immediately."
    )
    r_act.font.name = "Arial"
    r_act.font.size = Pt(10.5)
    r_act.font.color.rgb = DARK_TEXT
    p_act.paragraph_format.line_spacing = 1.15
    p_act.paragraph_format.space_after = Pt(14)

    # Sign-off
    p_sign = doc.add_paragraph()
    r_sign = p_sign.add_run("Sincerely,\n")
    r_sign.font.name = "Arial"
    r_sign.font.size = Pt(10)
    
    r_tf = p_sign.add_run("AI Task Force, COER University")
    r_tf.font.name = "Arial"
    r_tf.font.size = Pt(10.5)
    r_tf.font.bold = True
    r_tf.font.color.rgb = NAVY

    output_path = "/home/dell/Face-Attendance-System-Web-Version/AI_Task_Force_Official_Representation.docx"
    doc.save(output_path)
    print(f"Short document successfully updated at: {output_path}")

if __name__ == "__main__":
    create_short_document()
